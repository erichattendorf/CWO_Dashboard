import streamlit as st
import streamlit.components.v1 as components
import requests
import re
from datetime import datetime, timezone, timedelta
from zoneinfo import ZoneInfo
import os
import json
import pandas as pd
import io
import calendar
import docx
import holidays 

# --- PAGE SETUP ---
st.set_page_config(page_title="BHM CWO Dashboard", layout="wide", initial_sidebar_state="expanded")

# --- CSS HACK: TURN THE SIDEBAR ARROW INTO A GIANT "ALERTS" BUTTON ---
st.markdown("""
<style>
    [data-testid="collapsedControl"] {
        background-color: #ff4b4b !important;
        color: white !important;
        border-radius: 0 8px 8px 0 !important;
        padding: 5px 15px 5px 10px !important;
        box-shadow: 2px 2px 10px rgba(0,0,0,0.3) !important;
        width: auto !important;
        transition: 0.2s !important;
    }
    [data-testid="collapsedControl"]::after {
        content: " 🚨 ALERTS";
        font-family: sans-serif;
        font-weight: bold;
        font-size: 16px;
        margin-left: 5px;
    }
    [data-testid="collapsedControl"]:hover {
        background-color: #cc0000 !important;
    }
    [data-testid="collapsedControl"] title {
        display: none !important;
    }
</style>
""", unsafe_allow_html=True)

# --- DETERMINE CURRENT SHIFT (CENTRAL TIME) ---
bhm_tz = ZoneInfo("America/Chicago")
now_ct = datetime.now(bhm_tz)
hour = now_ct.hour

if 6 <= hour < 14:
    shift_name = "0600 - 1400 (Morning)"
    shift_date = now_ct.strftime("%Y-%m-%d")
elif 14 <= hour < 22:
    shift_name = "1400 - 2200 (Evening)"
    shift_date = now_ct.strftime("%Y-%m-%d")
else:
    shift_name = "2200 - 0600 (Night)"
    shift_date = (now_ct - timedelta(days=1)).strftime("%Y-%m-%d") if hour < 6 else now_ct.strftime("%Y-%m-%d")

current_shift_id = f"{shift_date}_{shift_name}"

# --- FILE DATABASE HELPER FUNCTIONS ---
LOGS_FILE = "shift_logs.json"
LEAVE_FILE = "leave_requests.json"
SCHED_CONFIG_FILE = "sched_config.json"

def load_json_db(filepath, default_val=[]):
    if os.path.exists(filepath):
        try:
            with open(filepath, "r") as f: return json.load(f)
        except: return default_val
    return default_val

def save_json_db(filepath, data):
    with open(filepath, "w") as f: json.dump(data, f, indent=4)

def parse_docx_to_df(file_bytes):
    try:
        doc = docx.Document(file_bytes)
        if len(doc.tables) > 0:
            table = doc.tables[0]
            data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if data: return pd.DataFrame(data[1:], columns=data[0])
                
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_lines = []
        in_table = False
        for line in lines:
            if line.startswith("DAY,DATE"): in_table = True
            if in_table:
                if "2200-0600" in line or line.startswith("M:"): break
                table_lines.append(line)
        
        if table_lines: return pd.read_csv(io.StringIO("\n".join(table_lines)))
        return None
    except Exception as e:
        st.error(f"Failed to read DOCX: {e}")
        return None

# --- INITIAL DEFAULT SCHEDULE SETUP ---
baseline_csv_string = """DAY,DATE,SP,RWB,TH,JA,MG,EJH,JDM,TRH
WED,1,M,D,E,,,,,OFF
THU,2,M ,D,E,,,,,-
FRI,3,M,D,E,A,,,,-
SAT,4,M,,,E,,D,,-
SUN,5,,,,E,M,,D,-
MON ,6,,D,E,,M,,,OFF
TUE,7,M,D,E,OFF,,,,
WED ,8,M,D,E,,,,,
THU,9,M,D,E,,,,,A
FRI,10,M,D,E,,,,,
SAT ,11,M,,,D,,,,E
SUN,12,,,,,M,,D,E
MON,13,,D,E,,M,,,
TUE,14,M,D,E,,,,,
WED,15,M ,D,SLV,,,,,E
THU,16,M,D,E,,,,,
FRI ,17,M,D,E,A,,,,OFF
SAT,18,M,,,E,,D,,OFF
SUN,19,,,,E,M,,D,OFF
MON,20,,D,E,,M,,,
TUE ,21,M,D,E,,,,,
WED ,22,M,D,E,,,,,
THU,23,M,D,E,,,,,A
FRI ,24,M,D,E,,,,,
SAT,25,M,,,D,,,,E
SUN,26,,,,,M,,D,E
MON,27,,D,E,,M,,,
TUE,28,M,D,E,,,,,
WED,29,M,D,E,A,,,,
THU,30,M,D,E,,,,,"""

# Dynamic Schedule Dictionary
sched_config = load_json_db(SCHED_CONFIG_FILE, default_val={"months": ["APRIL 2026"]})
if "APRIL 2026" not in sched_config["months"]: 
    sched_config["months"].insert(0, "APRIL 2026")
    save_json_db(SCHED_CONFIG_FILE, sched_config)

if not os.path.exists("baseline_APRIL_2026.csv"):
    with open("baseline_APRIL_2026.csv", "w") as f: f.write(baseline_csv_string)


# --- SIDEBAR: SHIFT DUTIES & ALERTS ---
with st.sidebar:
    st.title("📋 Shift Duties")
    st.caption(f"Current Shift: **{shift_name}**")
    
    checklist_items = [
        "Sign-on OID", "Sign-in Contractor & FAA Logs", "Receive Shift Briefing",
        "Check Digital T/Td Unit (Charged)", "Accomplish METAR/SPECI", "Check Read File",
        "Mid-shift: Time-check (Naval Obs)", "Mid-shift: Phone check (ATC/TRACON)",
        "QA TWO previous shifts", "Review OBS for electronic PMR"
    ]
    
    for item in checklist_items:
        st.checkbox(item, key=f"chk_{current_shift_id}_{item}")
        
    st.divider()
    
    st.title("⏰ METAR Alerts")
    st.markdown("Audio/Visual warnings before the XX:53 observation.")
    
    alarm_enabled = st.toggle("Enable Alerts", value=False)
    
    if alarm_enabled:
        alarm_minute = st.number_input("Trigger at XX past hour:", min_value=0, max_value=59, value=48, step=1)
        sound_choices = [
            "1. SOS Morse Code", "2. Classic Triple Beep", "3. Two-Tone Warning", 
            "4. Gentle Chimes", "5. Submarine Ping", "6. Harsh Buzzer", 
            "7. Digital Watch", "8. Air Raid Siren", "9. Radar Sweep", "10. Urgent Trill",
            "11. Telephone Ring", "12. Sci-Fi Alarm", "13. EKG Heart Monitor", 
            "14. Fast Geiger", "15. Deep Foghorn"
        ]
        alarm_sound = st.selectbox("Sound Profile:", sound_choices)
        sound_id = int(alarm_sound.split(".")[0])
        
        col_v1, col_v2 = st.columns(2)
        with col_v1: alarm_vol = st.slider("Vol %", min_value=1, max_value=100, value=50, step=1)
        with col_v2: alarm_pitch = st.slider("Pitch %", min_value=50, max_value=200, value=100, step=10)

        alarm_html = f"""
        <div id="idle-box" style="text-align:center; border-radius: 10px; padding: 15px; margin-top: 10px; background: #f0f2f6; border: 1px solid #ddd;">
            <h3 style="color: #333; margin: 0 0 10px 0; font-size: 16px;">Monitoring for XX:{alarm_minute:02d}</h3>
            <button onclick="triggerAlarmUI(true)" style="padding: 8px 15px; border-radius: 5px; border: 1px solid #aaa; cursor: pointer; background: white; font-weight: bold; color: #333;">🔊 Test Alert</button>
        </div>
        <div id="alert-box" style="display:none; background-color: #ff4b4b; color: white; padding: 20px; border-radius: 10px; text-align: center; margin-top: 10px;">
            <h2 style="margin: 0; font-size: 24px; font-family: sans-serif;">🚨 ALARM 🚨</h2>
            <h3 style="margin: 5px 0 15px 0; font-size: 16px; font-weight: normal;">Observation Due!</h3>
            <button onclick="silenceAlarm()" style="padding: 15px; font-size: 16px; font-weight: bold; color: #ff4b4b; background: white; border: none; border-radius: 5px; cursor: pointer; width: 100%;">
                🔕 Silence Current Alarm
            </button>
        </div>
        <script>
        const AudioContext = window.AudioContext || window.webkitAudioContext;
        let audioCtx;
        let activeOscillators = [];
        let alarmInterval;
        let isSilencedForThisMinute = false;

        function playTone(freq, type, startDelay, duration, vol) {{
            const pitchMult = {alarm_pitch} / 100.0;
            const finalFreq = freq * pitchMult;
            const osc = audioCtx.createOscillator();
            const gain = audioCtx.createGain();
            osc.type = type;
            osc.frequency.value = finalFreq;
            osc.connect(gain);
            gain.connect(audioCtx.destination);
            const t = audioCtx.currentTime + startDelay;
            gain.gain.setValueAtTime(0, t);
            gain.gain.linearRampToValueAtTime(vol, t + 0.02);
            gain.gain.setValueAtTime(vol, t + duration - 0.02);
            gain.gain.linearRampToValueAtTime(0, t + duration);
            osc.start(t);
            osc.stop(t + duration);
            activeOscillators.push(osc);
        }}

        function playAlarmAudio() {{
            if (!audioCtx) audioCtx = new AudioContext();
            if (audioCtx.state === 'suspended') audioCtx.resume();
            activeOscillators.forEach(osc => {{ try {{ osc.stop(); }} catch(e){{}} }});
            activeOscillators = [];
            const rawVol = {alarm_vol} / 100.0;
            const v = Math.pow(rawVol, 2);
            const s = {sound_id};

            if (s === 1) {{ 
                let t = 0;
                for(let i=0; i<3; i++) {{
                    playTone(700, 'sine', t, 0.1, v); playTone(700, 'sine', t+0.2, 0.1, v); playTone(700, 'sine', t+0.4, 0.1, v); t += 0.8;
                    playTone(700, 'sine', t, 0.3, v); playTone(700, 'sine', t+0.4, 0.3, v); playTone(700, 'sine', t+0.8, 0.3, v); t += 1.3;
                    playTone(700, 'sine', t, 0.1, v); playTone(700, 'sine', t+0.2, 0.1, v); playTone(700, 'sine', t+0.4, 0.1, v); t += 1.5;
                }}
            }} 
            else if (s === 2) {{ playTone(800, 'sine', 0, 0.15, v); playTone(800, 'sine', 0.3, 0.15, v); playTone(800, 'sine', 0.6, 0.15, v); }}
            else if (s === 3) {{ for(let i=0; i<4; i++) {{ playTone(600, 'square', i*1.0, 0.5, v*0.5); playTone(800, 'square', (i*1.0)+0.5, 0.5, v*0.5); }} }}
            else if (s === 4) {{ playTone(523.25, 'sine', 0, 0.4, v); playTone(659.25, 'sine', 0.3, 0.4, v); playTone(783.99, 'sine', 0.6, 0.8, v); }}
            else if (s === 5) {{ playTone(1000, 'sine', 0, 0.1, v); playTone(1000, 'sine', 0.1, 1.5, v*0.1); }}
            else if (s === 6) {{ playTone(150, 'sawtooth', 0, 0.5, v); playTone(150, 'sawtooth', 0.8, 0.5, v); playTone(150, 'sawtooth', 1.6, 0.5, v); }}
            else if (s === 7) {{ playTone(2000, 'square', 0, 0.08, v*0.2); playTone(2000, 'square', 0.15, 0.08, v*0.2); playTone(2000, 'square', 1.0, 0.08, v*0.2); playTone(2000, 'square', 1.15, 0.08, v*0.2); }}
            else if (s === 8) {{
                const osc = audioCtx.createOscillator(); const gain = audioCtx.createGain();
                osc.type = 'sine'; osc.connect(gain); gain.connect(audioCtx.destination);
                const pitchMult = {alarm_pitch} / 100.0;
                gain.gain.setValueAtTime(0, audioCtx.currentTime); gain.gain.linearRampToValueAtTime(v, audioCtx.currentTime + 0.5);
                gain.gain.setValueAtTime(v, audioCtx.currentTime + 3.5); gain.gain.linearRampToValueAtTime(0, audioCtx.currentTime + 4.0);
                osc.frequency.setValueAtTime(400 * pitchMult, audioCtx.currentTime); osc.frequency.linearRampToValueAtTime(1200 * pitchMult, audioCtx.currentTime + 2);
                osc.frequency.linearRampToValueAtTime(400 * pitchMult, audioCtx.currentTime + 4);
                osc.start(audioCtx.currentTime); osc.stop(audioCtx.currentTime + 4.0); activeOscillators.push(osc);
            }}
            else if (s === 9) {{ for(let i=0; i<3; i++) {{ playTone(900, 'triangle', i*1.5, 0.1, v); }} }}
            else if (s === 10) {{ for(let i=0; i<15; i++) {{ playTone(i%2==0 ? 700 : 750, 'square', i*0.1, 0.08, v*0.5); }} }}
            else if (s === 11) {{ for(let i=0; i<6; i++) {{ playTone(1200, 'sine', i*0.15, 0.05, v); playTone(1500, 'sine', i*0.15+0.05, 0.05, v); }} }}
            else if (s === 12) {{
                for(let i=0; i<5; i++) {{
                    const osc = audioCtx.createOscillator(); const gain = audioCtx.createGain();
                    osc.type = 'sawtooth'; osc.connect(gain); gain.connect(audioCtx.destination);
                    const pitchMult = {alarm_pitch} / 100.0;
                    gain.gain.setValueAtTime(v, audioCtx.currentTime + (i*0.5)); gain.gain.linearRampToValueAtTime(0, audioCtx.currentTime + (i*0.5) + 0.4);
                    osc.frequency.setValueAtTime(1500 * pitchMult, audioCtx.currentTime + (i*0.5)); osc.frequency.linearRampToValueAtTime(300 * pitchMult, audioCtx.currentTime + (i*0.5) + 0.4);
                    osc.start(audioCtx.currentTime + (i*0.5)); osc.stop(audioCtx.currentTime + (i*0.5) + 0.4); activeOscillators.push(osc);
                }}
            }}
            else if (s === 13) {{ playTone(900, 'sine', 0, 0.1, v); playTone(900, 'sine', 1.0, 0.1, v); playTone(900, 'sine', 2.0, 0.1, v); }}
            else if (s === 14) {{ for(let i=0; i<20; i++) {{ playTone(2000, 'square', i*0.1 + (Math.random()*0.05), 0.02, v*0.2); }} }}
            else if (s === 15) {{ playTone(100, 'sawtooth', 0, 2.0, v); playTone(102, 'square', 0, 2.0, v*0.5); }}
        }}

        window.silenceAlarm = function() {{
            clearInterval(alarmInterval);
            activeOscillators.forEach(osc => {{ try {{ osc.stop(); }} catch(e){{}} }});
            activeOscillators = [];
            document.getElementById('alert-box').style.display = 'none';
            document.getElementById('idle-box').style.display = 'block';
            isSilencedForThisMinute = true;
        }};

        function triggerAlarmUI(isTest = false) {{
            try {{
                const parent = window.parent.document;
                const expandBtn = parent.querySelector('[data-testid="collapsedControl"]');
                if (expandBtn && expandBtn.getAttribute('aria-expanded') !== 'true') {{ expandBtn.click(); }}
            }} catch(e) {{}}
            document.getElementById('idle-box').style.display = 'none';
            document.getElementById('alert-box').style.display = 'block';
            playAlarmAudio();
            if (!isTest) {{ alarmInterval = setInterval(playAlarmAudio, 5000); }}
        }}

        let hasTriggeredThisHour = false;
        setInterval(function() {{
            var d = new Date();
            if (d.getMinutes() === {alarm_minute}) {{
                if (!hasTriggeredThisHour && !isSilencedForThisMinute) {{
                    triggerAlarmUI(false);
                    hasTriggeredThisHour = true;
                }}
            }} else {{
                hasTriggeredThisHour = false; 
                isSilencedForThisMinute = false;
                document.getElementById('alert-box').style.display = 'none';
                document.getElementById('idle-box').style.display = 'block';
            }}
        }}, 1000);
        </script>
        """
        components.html(alarm_html, height=160)
    else:
        st.info("🔕 Alerts Disabled.")

# --- HEADER WITH LOGOS ---
header_col1, header_col2 = st.columns([4, 2])

with header_col1:
    st.title("BHM CWO Tactical Dashboard 🌪️")
    st.subheader("Logic Engine & Live Interface")

with header_col2:
    st.markdown("<br>", unsafe_allow_html=True)
    logo1, logo2 = st.columns(2)
    with logo1:
        if os.path.exists("Cat and Hat.jpg"):
            st.image("Cat and Hat.jpg", width=250)
            st.caption("**Created by Eric Hattendorf**")
    with logo2:
        if os.path.exists("NWS.png"):
            st.image("NWS.png", width=100)

st.divider()

# --- FUNCTIONS ---
NWS_HEADERS = {"User-Agent": "BHM_CWO_Dashboard/2.0", "Accept": "application/geo+json"}

def parse_nws_properties(props):
    timestamp = props.get('timestamp')
    if not timestamp: return None, None, None, None
    dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
    time_str = dt.strftime("%H:%MZ")
    
    try:
        temp_c = float(props.get('temperature', {}).get('value'))
        dew_c = float(props.get('dewpoint', {}).get('value'))
    except (TypeError, ValueError): temp_c, dew_c = None, None

    wdir = props.get('windDirection', {}).get('value')
    wspd_kmh = props.get('windSpeed', {}).get('value')
    wdir_str = f"{int(wdir):03d}" if wdir is not None else "VRB"
    wspd_kts = int(round(wspd_kmh / 1.852)) if wspd_kmh is not None else 0
    wind_str = f"{wdir_str}@{wspd_kts}kt"
    
    vis_m = props.get('visibility', {}).get('value')
    vis_sm = round(vis_m / 1609.34, 1) if vis_m is not None else "M"
    
    clouds_str = ""
    cloud_layers = props.get('cloudLayers') or []
    if cloud_layers:
        layer_strs = []
        for layer in cloud_layers:
            amt = layer.get('amount', '---')
            base_dict = layer.get('base')
            base_m = base_dict.get('value') if isinstance(base_dict, dict) else None
            if base_m is not None and amt not in ["CLR", "SKC"]:
                base_ft = base_m * 3.28084
                base_hnds = int(round(base_ft / 100))
                layer_strs.append(f"{amt}{base_hnds:03d}")
            else:
                layer_strs.append(amt)
        clouds_str = " ".join(layer_strs)
        
    if not clouds_str:
        desc = props.get('textDescription', '')
        if desc and "Clear" in desc: clouds_str = "CLR"
        else: clouds_str = "M"
            
    temp_str = f"{int(temp_c)}" if temp_c is not None else "M"
    dew_str = f"{int(dew_c)}" if dew_c is not None else "M"
    pres_pa = props.get('barometricPressure', {}).get('value')
    pres_inhg = round(pres_pa / 3386.389, 2) if pres_pa is not None else "M"
    
    formatted_str = f"({time_str}) | Wind: {wind_str} | Vis: {vis_sm}SM | Sky: {clouds_str} | Temp/Dew: {temp_str}/{dew_str}°C | Alt: {pres_inhg} inHg"
    return formatted_str, timestamp, temp_c, dew_c

def get_5min_asos():
    try:
        url = "https://api.weather.gov/stations/KBHM/observations?limit=10"
        response = requests.get(url, headers=NWS_HEADERS, timeout=10)
        if response.status_code == 200:
            features = response.json().get('features', [])
            if len(features) > 0:
                ob = features[0]
                props = ob['properties']
                raw = props.get('rawMessage')
                formatted_str, timestamp, temp_c, dew_c = parse_nws_properties(props)
                if formatted_str:
                    final_raw = f"BHM 5-MIN ({timestamp[11:16]}Z) | {raw}" if raw else f"BHM 5-MIN {formatted_str}"
                    return final_raw, timestamp, temp_c, dew_c, None
            return None, None, None, None, "No data features found in NWS ping."
        else: return None, None, None, None, f"HTTP {response.status_code}"
    except Exception as e: return None, None, None, None, f"Error: {str(e)}"

def get_regional_5min():
    stations = ["KTCL", "KANB", "KEET", "KPLR"]
    data = {}
    for stn in stations:
        try:
            url = f"https://api.weather.gov/stations/{stn}/observations?limit=6"
            res = requests.get(url, headers=NWS_HEADERS, timeout=5)
            if res.status_code == 200:
                features = res.json().get('features', [])
                parsed = []
                for ob in features:
                    props = ob['properties']
                    raw = props.get('rawMessage')
                    formatted_str, ts, _, _ = parse_nws_properties(props)
                    if formatted_str:
                        parsed.append(f"({ts[11:16]}Z) | {raw}" if raw else formatted_str)
                data[stn] = parsed
            else: data[stn] = [f"API Error: HTTP {res.status_code}"]
        except: data[stn] = ["Connection Error"]
    return data

def get_awc_data():
    try:
        url = "https://aviationweather.gov/api/data/metar?ids=KBHM&format=json&hours=6"
        response = requests.get(url, headers=NWS_HEADERS, timeout=10)
        if response.status_code == 200: 
            data = response.json()
            if isinstance(data, list) and len(data) > 0:
                return data
        return None
    except: return None

def get_taf_data():
    try:
        url = "https://aviationweather.gov/api/data/taf?ids=KBHM&format=json"
        response = requests.get(url, headers=NWS_HEADERS, timeout=10)
        if response.status_code == 200:
            data = response.json()
            if isinstance(data, list) and len(data) > 0:
                return data[0].get('rawTAF', 'TAF currently unavailable.')
        return "TAF currently unavailable."
    except: return "Error fetching TAF."

def extract_vis_and_cig(metar_string):
    vis_val, cig_val = 10.0, 10000 
    if "Error" in metar_string or "No recent" in metar_string or not metar_string: return vis_val, cig_val
    vis_match = re.search(r'\s(M?P?\d+/?\d*\s?\d*/?\d*)SM', metar_string)
    if vis_match:
        vis_str = vis_match.group(1).replace('P', '').replace('M', '').strip()
        try:
            if " " in vis_str:
                whole, frac = vis_str.split(" ")
                num, den = frac.split("/")
                vis_val = float(whole) + (float(num) / float(den))
            elif "/" in vis_str:
                num, den = vis_str.split("/")
                vis_val = float(num) / float(den)
            else: vis_val = float(vis_str)
        except: pass 
    cig_match = re.search(r'(BKN|OVC|VV)(\d{3})', metar_string)
    if cig_match: cig_val = int(cig_match.group(2)) * 100
    return vis_val, cig_val

def check_speci(old_val, new_val, thresholds):
    for t in thresholds:
        if (old_val >= t and new_val < t) or (old_val < t and new_val >= t): return True, t
    return False, None

vis_thresholds = [3.0, 2.0, 1.0, 0.5, 0.25]
cig_thresholds = [3000, 1500, 1000, 500, 200]

# --- COMM CHECK UI ---
st.markdown("### 📡 System Comm Status (5-Min Early Warning)")
latest_raw, latest_ts, live_temp_c, live_dew_c, diag_err = get_5min_asos()

if latest_raw and latest_ts:
    try:
        ob_time = datetime.fromisoformat(latest_ts.replace('Z', '+00:00'))
        now = datetime.now(timezone.utc)
        age_minutes = (now - ob_time).total_seconds() / 60
        col_a, col_b = st.columns([2, 1])
        with col_a: st.code(f"{latest_raw}", language="bash")
        with col_b:
            if age_minutes > 45: st.error(f"🚨 **COMM WARNING:** Ping is **{int(age_minutes)} mins old!** Check long-line.")
            else: st.success(f"✅ **Comms Good:** Latency is **{int(age_minutes)}** mins.")
    except: st.warning("Could not calculate age.")
else: st.warning(f"⚠️ 5-minute network ping unavailable. Diagnostic code: {diag_err}")

st.divider()

# --- TOP UI: OBSERVATIONS & RADAR ---
top_col1, top_col2 = st.columns([1, 1])

with top_col1:
    tab_local, tab_regional = st.tabs(["📍 KBHM Transmitted", "🌍 Regional 5-Min (50-mi)"])
    
    with tab_local:
        awc_data = get_awc_data()
        recent_metars = [obs.get('rawOb', 'No raw string') for obs in awc_data[:6]] if awc_data else ["No recent METARs found."]
        latest_metar = recent_metars[0] if len(recent_metars) > 0 else ""
        live_vis, live_cig = extract_vis_and_cig(latest_metar)
        
        if "No recent" not in latest_metar:
            for i, metar in enumerate(recent_metars):
                if i == 0: st.error(f"**LATEST:** `{metar}`")
                elif i == 1: st.warning(f"`{metar}`")
                else: st.info(f"`{metar}`")
        else: st.warning("Could not load AWC METARs.")
        
        st.markdown("---")
        st.markdown("#### ✈️ KBHM Current TAF")
        taf_string = get_taf_data()
        if taf_string != "TAF currently unavailable." and taf_string != "Error fetching TAF.":
            formatted_taf = re.sub(r'(FM\d{6}|TEMPO|BECMG|PROB)', r'\n\1', taf_string)
            st.info(formatted_taf)
        else:
            st.warning(taf_string)

    with tab_regional:
        st.markdown("**Last Hour of 5-Minute ASOS Data**")
        if st.button("Fetch Fresh Regional Data"):
            with st.spinner("Pulling regional observations..."):
                reg_data = get_regional_5min()
                for stn, obs_list in reg_data.items():
                    st.markdown(f"**{stn}**")
                    for ob in obs_list: st.caption(f"`{ob}`")

with top_col2:
    st.markdown("#### 📡 Auto-Updating Radar (KBMX)")
    radar_html = """
    <div style="text-align:center; height: 100%;">
        <img id="auto-radar" src="https://radar.weather.gov/ridge/standard/KBMX_loop.gif" style="width:100%; max-height: 550px; object-fit: contain; border-radius:10px;">
    </div>
    <script>
        setInterval(function() {
            var img = document.getElementById("auto-radar");
            var timestamp = new Date().getTime();
            img.src = "https://radar.weather.gov/ridge/standard/KBMX_loop.gif?t=" + timestamp;
        }, 300000);
    </script>
    """
    components.html(radar_html, height=580)

st.divider()

# --- MIDDLE UI: CALCULATORS, CONTACTS, LOGS ---
st.subheader("🧮 SOPs, Tools & Operations")
calc_tab, cont_tab, log_tab = st.tabs(["Calculators", "📞 Contacts & SOPs", "📖 Shift Logs"])

with calc_tab:
    calc_col1, calc_col2, calc_col3 = st.columns(3)
    with calc_col1:
        st.markdown("**🌫️ Visibility SPECI**")
        old_vis = st.number_input("Prev Vis (SM)", min_value=0.0, value=10.0, step=0.25)
        new_vis = st.number_input("Cur Vis (SM)", min_value=0.0, value=float(live_vis), step=0.25)
        vis_speci, vis_trigger = check_speci(old_vis, new_vis, vis_thresholds)
        if old_vis != new_vis:
            if vis_speci: st.error(f"🚨 **SPECI:** Crossed {vis_trigger} SM.")
            else: st.success("✅ No SPECI required.")
    
    with calc_col2:
        st.markdown("**☁️ Ceiling SPECI**")
        old_cig = st.number_input("Prev Ceiling (Ft)", min_value=0, value=5000, step=100)
        new_cig = st.number_input("Cur Ceiling (Ft)", min_value=0, value=int(live_cig), step=100)
        cig_speci, cig_trigger = check_speci(old_cig, new_cig, cig_thresholds)
        if old_cig != new_cig:
            if cig_speci: st.error(f"🚨 **SPECI:** Crossed {cig_trigger} FT.")
            else: st.success("✅ No SPECI required.")

    with calc_col3:
        st.markdown("**⚡ Flash-to-Bang (Lightning)**")
        st.caption("Rules: **0-5m** = OHD | **5-10m** = VC | **10-30m** = DSNT")
        sec_delay = st.number_input("Seconds between Flash and Thunder:", min_value=0, value=0, step=1)
        if sec_delay > 0:
            miles = round(sec_delay / 5.0, 1)
            if miles <= 5: st.error(f"**{miles} Miles (TS OHD)**")
            elif miles <= 10: st.warning(f"**{miles} Miles (TS VC)**")
            else: st.success(f"**{miles} Miles (TS DSNT)**")
        else:
            st.caption("Divide seconds by 5 to get distance in miles.")

with cont_tab:
    st.markdown("### 📞 BHM Emergency Contacts & IT")
    col_c1, col_c2, col_c3 = st.columns(3)
    with col_c1:
        st.markdown("**FAA / ATC**\n* BHM ATCT CAB: `205-769-3914`\n* BHM TRACON: `205-769-3907`\n* Tech Ops: `205-769-3950`\n* NWS Office: `205-621-5650`")
    with col_c2:
        st.markdown("**Maintenance / IT**\n* Jeff Short (WDI): `641-923-6043`\n* BHM IT (Free WiFi): `205-599-0700`\n* Facilities (Mathew): `205-595-0595`\n* BHM Voice: `205-591-6172`")
    with col_c3:
        st.markdown("**ASOS OUTAGES**\n* AOMC: `1-800-242-8194` or `8895`\n* NEMC (Comms): `1-855-322-6362` *(-> ATL -> #3)*\n* ARTCC Longline: `1-770-210-7960` *(Cell phone only)*")

with log_tab:
    st.markdown("### Persistent Shift Notes Database")
    logs = load_json_db(LOGS_FILE)
    
    new_log = st.text_area("Enter a new shift note/passdown:", height=100)
    col_l1, col_l2 = st.columns([1, 5])
    with col_l1:
        if st.button("💾 Save Log"):
            if new_log.strip() != "":
                entry = {
                    "timestamp": datetime.now(bhm_tz).strftime("%Y-%m-%d %H:%M CT"),
                    "shift": shift_name,
                    "note": new_log.strip()
                }
                logs.insert(0, entry) 
                save_json_db(LOGS_FILE, logs)
                st.success("Saved!")
                st.rerun()
            else: st.warning("Note is empty.")
    
    st.markdown("---")
    st.markdown("#### Passdown History")
    if logs:
        for l in logs: st.info(f"**{l['timestamp']}** ({l['shift']})\n\n{l['note']}")
    else: st.caption("No logs currently in database.")

st.divider()

# --- BOTTOM UI: REMARKS (RMK) BUILDER ---
st.subheader("📝 Remarks (RMK) Builder (A-Z Strict Order)")
tab_svr, tab_wind, tab_ts, tab_precip = st.tabs(["🌪️ Severe Wx", "💨 Wind/Pressure", "🌩️ TS & Lightning", "🌧️ Precip/Vis/Hail"])
rmks = {letter: "" for letter in "ABCDEFGHIJKLMNOPQRSTUVW"}

with tab_svr:
    col_a, col_b = st.columns(2)
    with col_a:
        has_torn = st.checkbox("Tornadic Activity")
        if has_torn:
            st.error("🚨 MAIN OBS REMINDER: Add `+FC` or `FC` to Present WX.")
            torn_type = st.selectbox("Type:", ["TORNADO", "FUNNEL CLOUD", "WATERSPOUT"])
            torn_b, torn_e = st.text_input("Begin Time (Min):", ""), st.text_input("End Time (Min):", "")
            torn_loc, torn_mov = st.selectbox("Location:", ["", "ALQDS", "N", "NE", "E", "SE", "S", "SW", "W", "NW"]), st.selectbox("Moving:", ["", "N", "NE", "E", "SE", "S", "SW", "W", "NW"])
            t_str = f"{torn_type}" + (f" B{torn_b}" if torn_b else "") + (f" E{torn_e}" if torn_e else "") + (f" {torn_loc}" if torn_loc else "") + (f" MOV {torn_mov}" if torn_mov else "")
            rmks['B'] = t_str
    with col_b:
        has_volc = st.checkbox("Volcanic Eruption")
        if has_volc:
            st.error("🚨 MAIN OBS REMINDER: Add `VA` to Present WX.")
            volc_name = st.text_input("Volcano Name:")
            if volc_name: rmks['A'] = f"MT {volc_name.upper()} ERUPTION"

with tab_wind:
    col_w1, col_w2, col_w3 = st.columns(3)
    with col_w1:
        has_pk_wnd = st.checkbox("Peak Wind (>25kt)")
        if has_pk_wnd:
            pk_dir, pk_spd, pk_time = st.number_input("Direction", min_value=0, max_value=360, step=10, value=270), st.number_input("Speed (Kts)", min_value=26, max_value=200, step=1, value=35), st.number_input("Time (Min)", min_value=0, max_value=59, step=1, value=15)
            rmks['C'] = f"PK WND {pk_dir:03d}{pk_spd}/{pk_time:02d}"
    with col_w2:
        has_wshft = st.checkbox("Wind Shift")
        if has_wshft:
            wshft_time = st.number_input("Shift Time (Min)", min_value=0, max_value=59, step=1, value=15)
            rmks['D'] = f"WSHFT {wshft_time:02d}" + (" FROPA" if st.checkbox("FROPA") else "")
    with col_w3:
        pressure_rmk = st.radio("Pressure Trend:", ["None", "Rising Rapidly", "Falling Rapidly"])
        if pressure_rmk == "Rising Rapidly": rmks['R'] = "PRESRR"
        elif pressure_rmk == "Falling Rapidly": rmks['R'] = "PRESFR"

with tab_ts:
    col_ts1, col_ts2 = st.columns(2)
    with col_ts1:
        has_ltg = st.checkbox("Lightning Observed")
        if has_ltg:
            ltg_freq = st.selectbox("Frequency:", ["OCNL", "FRQ", "CONS"])
            ltg_types = st.multiselect("Type:", ["IC", "CG", "CC", "CA"], default=["CG"])
            ltg_loc = st.selectbox("LTG Location:", ["ALQDS", "OHD", "VC", "DSNT", "N", "NE", "E", "SE", "S", "SW", "W", "NW"])
            rmks['H'] = f"{ltg_freq} LTG{''.join(ltg_types)} {ltg_loc}"
            
    with col_ts2:
        has_ts = st.checkbox("Thunderstorm Active")
        if has_ts:
            st.warning("🚨 **TS REMINDER:** Put `TS` or `VCTS` in Pres WX | Add `CB` to Sky | Turn ALDARS to `MAN`")
            
            ts_dist = st.selectbox("Distance Category:", ["Overhead (OHD, <5SM)", "Vicinity (VC, 5-10SM)", "Distant (DSNT, 10-30SM)"])
            ts_dir = st.selectbox("Direction:", ["", "ALQDS", "N", "NE", "E", "SE", "S", "SW", "W", "NW"])
            ts_mov = st.selectbox("TS Moving:", ["Unknown", "N", "NE", "E", "SE", "S", "SW", "W", "NW"])
            
            # Formats precisely to JO 7900.5E criteria
            if ts_dist.startswith("Overhead"):
                ts_str = "TS OHD"
            elif ts_dist.startswith("Vicinity"):
                ts_str = f"TS {ts_dir}".strip() if ts_dir else "TS"
            elif ts_dist.startswith("Distant"):
                ts_str = f"TS DSNT {ts_dir}".strip() if ts_dir else "TS DSNT"
                
            rmks['K'] = ts_str + (f" MOV {ts_mov}" if ts_mov != "Unknown" else "")

with tab_precip:
    col_p1, col_p2, col_p3, col_p4 = st.columns(4)
    with col_p1:
        st.markdown("**🌧️ Precipitation**")
        has_precip = st.checkbox("Precip Begin/End")
        if has_precip:
            p_type, p_b, p_e = st.selectbox("Precip Type:", ["RA", "SN", "DZ", "UP"]), st.text_input("B (Min past HR):", ""), st.text_input("E (Min past HR):", "")
            rmks['I'] = f"{p_type}" + (f"B{p_b}" if p_b else "") + (f"E{p_e}" if p_e else "")
    with col_p2:
        st.markdown("**🧊 Hail**")
        has_hail = st.checkbox("Hail (GR/GS)")
        if has_hail:
            st.error("🚨 REMINDER: `GR` (>=1/4 in) or `GS` (<1/4 in).")
            rmks['L'] = f"GR {st.text_input('Hail Size:', '1/4')}"
    with col_p3:
        st.markdown("**🌫️ Virga**")
        has_virga = st.checkbox("VIRGA")
        if has_virga:
            virga_loc = st.selectbox("Virga Loc:", ["ALQDS", "N", "NE", "E", "SE", "S", "SW", "W", "NW"])
            rmks['M'] = "VIRGA ALQDS" if virga_loc == "ALQDS" else f"VIRGA {virga_loc}"
    with col_p4:
        st.markdown("**👁️ Visibility Remarks**")
        has_sector_vis = st.checkbox("Sector Vis (RMK G)")
        if has_sector_vis:
            sec_dir = st.selectbox("Sector:", ["N", "NE", "E", "SE", "S", "SW", "W", "NW"])
            sec_vis = st.text_input("Vis (e.g. 1 1/2):", "1")
            rmks['G'] = f"VIS {sec_dir} {sec_vis}"
        has_vrb_vis = st.checkbox("Variable Vis (RMK F)")
        if has_vrb_vis:
            vrb_min, vrb_max = st.text_input("Min Vis:", "1/2"), st.text_input("Max Vis:", "2")
            rmks['F'] = f"VIS {vrb_min}V{vrb_max}"

st.markdown("---")
final_remarks = [rmks[key] for key in sorted(rmks.keys()) if rmks[key] != ""]
if final_remarks:
    st.success("**Final JO 7900.5E Remark String (Correct Order):**")
    st.code("RMK " + " ".join(final_remarks), language="markdown")

st.divider()

# --- ADMIN TOOLS: SCHEDULE & LEAVE ---
st.subheader("📅 Admin: Scheduling & Operations")
sched_tab, leave_tab = st.tabs(["Monthly Schedule Library & Payroll", "🏖️ 12-Month Leave Calendar"])

with sched_tab:
    col_t1, col_t2 = st.columns([1, 1])
    with col_t1:
        selected_month = st.selectbox("Select Schedule to View/Edit:", sched_config["months"])
    with col_t2:
        st.info("Upload a completely new month to the library here. (CSV or DOCX)")
        new_month_name = st.text_input("New Month Name (e.g., JUNE 2026):")
        uploaded_file = st.file_uploader("Upload Schedule File", type=['csv', 'docx'])
        if st.button("➕ Add New Month to Library"):
            if uploaded_file and new_month_name:
                df_new = None
                if uploaded_file.name.endswith('.csv'): df_new = pd.read_csv(uploaded_file)
                elif uploaded_file.name.endswith('.docx'): df_new = parse_docx_to_df(uploaded_file)
                
                if df_new is not None:
                    safe_name = new_month_name.replace(" ", "_")
                    df_new.to_csv(f"baseline_{safe_name}.csv", index=False)
                    if new_month_name not in sched_config["months"]:
                        sched_config["months"].append(new_month_name)
                        save_json_db(SCHED_CONFIG_FILE, sched_config)
                    st.success(f"{new_month_name} added to library!")
                    st.rerun()
                else: st.error("Could not extract a readable table.")
            else: st.warning("Please provide a name and upload a file.")

    st.markdown("---")
    safe_selected = selected_month.replace(" ", "_")
    base_file = f"baseline_{safe_selected}.csv"
    edit_file = f"edited_{safe_selected}.csv"
    
    if not os.path.exists(base_file):
        st.warning(f"File for {selected_month} is missing.")
    else:
        df_base = pd.read_csv(base_file)
        if os.path.exists(edit_file): df_curr = pd.read_csv(edit_file)
        else: df_curr = df_base.copy()
            
        st.markdown(f"**BIRMINGHAM AL CWO {selected_month} SCHEDULE**")
        st.caption("Legend: M: 2200-0600 | A: 0600-1400 | D: 0600-1400 | E: 1400-2200 | LV: Leave | H: Holiday | - : Off")
        
        edited_df = st.data_editor(df_curr, num_rows="dynamic", use_container_width=True, key=f"editor_{selected_month}")
        
        if st.button(f"💾 Save {selected_month} Edits"):
            edited_df.to_csv(edit_file, index=False)
            st.success("Schedule Updated!")
            st.rerun()
            
        st.markdown("---")
        st.markdown(f"### 🔴 {selected_month} Redline Tracker")
        st.caption("Shows changes made from the original published baseline schedule.")
        
        html = "<table style='width:100%; border-collapse: collapse; text-align: center; font-size: 14px;'>"
        html += "<tr style='background-color: #f0f2f6;'>" + "".join([f"<th style='border: 1px solid #ddd; padding: 8px; color: #333;'>{c}</th>" for c in df_base.columns]) + "</tr>"
        for i in range(len(df_base)):
            html += "<tr>"
            for col in df_base.columns:
                val_base = str(df_base.loc[i, col])
                try: val_curr = str(df_curr.loc[i, col])
                except KeyError: val_curr = "NaN"
                    
                if pd.isna(df_base.loc[i, col]): val_base = ""
                if i < len(df_curr) and pd.isna(df_curr.loc[i, col]): val_curr = ""
                if val_base == "nan": val_base = ""
                if val_curr == "nan": val_curr = ""

                if val_base != val_curr: cell_html = f"<del style='color:red;'>{val_base}</del><br><span style='color:red; font-weight:bold;'>{val_curr}</span>"
                else: cell_html = f"<span style='color: #333;'>{val_curr}</span>"
                html += f"<td style='border: 1px solid #ddd; padding: 8px;'>{cell_html}</td>"
            html += "</tr>"
        html += "</table>"
        st.markdown(html, unsafe_allow_html=True)
        
        # --- AUTOMATED PAYROLL GENERATOR WITH PAY PERIODS ---
        st.markdown("---")
        st.markdown("### ⏱️ Automated Payroll & Timesheet Generator")
        
        pay_period = st.radio("Select Pay Period:", ["Full Month", "1st - 15th", "16th - End of Month"], horizontal=True)
        
        if st.button(f"Generate Timesheet for {selected_month}"):
            payroll_data = []
            observers = df_curr.columns[2:] 
            
            for person in observers:
                total_hrs, nsd_hrs, sun_hrs = 0, 0, 0
                for index, row in df_curr.iterrows():
                    try: day_num = int(row['DATE'])
                    except: continue 

                    if pay_period == "1st - 15th" and day_num > 15: continue
                    if pay_period == "16th - End of Month" and day_num <= 15: continue

                    shift = str(row[person]).strip().upper()
                    day = str(row['DAY']).strip().upper()
                    
                    if shift in ['M', 'A', 'D', 'E']:
                        total_hrs += 8
                        if day == 'SUN': sun_hrs += 8
                        if shift == 'M': nsd_hrs += 8
                        elif shift == 'E': nsd_hrs += 4
                            
                payroll_data.append({
                    "Observer": person, 
                    "Total Hours": total_hrs, 
                    "Sunday Premium": sun_hrs, 
                    "Night Differential": nsd_hrs
                })
                
            df_payroll = pd.DataFrame(payroll_data)
            st.dataframe(df_payroll, use_container_width=True)
            
            csv_payroll = df_payroll.to_csv(index=False)
            st.download_button(
                label=f"⬇️ Download {pay_period} Payroll CSV",
                data=csv_payroll,
                file_name=f"Payroll_{safe_selected}_{pay_period.replace(' ', '')}.csv",
                mime="text/csv"
            )

with leave_tab:
    st.markdown("### 🏖️ Employee Time-Off Requests")
    leave_requests = load_json_db(LEAVE_FILE)
    
    col_req1, col_req2, col_req3, col_req4 = st.columns(4)
    with col_req1: req_name = st.text_input("Observer Name")
    with col_req2: req_start = st.date_input("Start Date")
    with col_req3: req_end = st.date_input("End Date")
    with col_req4: 
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("➕ Submit Request"):
            if req_name:
                entry = {
                    "submitted": datetime.now(bhm_tz).strftime("%Y-%m-%d"),
                    "name": req_name,
                    "start": req_start.strftime("%Y-%m-%d"),
                    "end": req_end.strftime("%Y-%m-%d")
                }
                leave_requests.append(entry)
                save_json_db(LEAVE_FILE, leave_requests)
                st.success("Request logged!")
                st.rerun()
            else: st.warning("Please enter a name.")

    # --- CUSTOM 12-MONTH VISUAL CALENDAR GENERATOR (WITH HOLIDAYS & TEXT COLOR FIX) ---
    st.markdown("---")
    st.markdown("#### 12-Month Visual Leave Calendar")
    
    leave_dict = {}
    for req in leave_requests:
        try:
            start_date = datetime.strptime(req['start'], "%Y-%m-%d")
            end_date = datetime.strptime(req['end'], "%Y-%m-%d")
            delta = end_date - start_date
            for i in range(delta.days + 1):
                d_str = (start_date + timedelta(days=i)).strftime("%Y-%m-%d")
                if d_str not in leave_dict: leave_dict[d_str] = []
                leave_dict[d_str].append(req['name'])
        except Exception as e: pass

    cur_year = datetime.now(bhm_tz).year
    cur_month = datetime.now(bhm_tz).month
    us_holidays = holidays.US(years=[cur_year, cur_year+1])

    class LeaveCalendar(calendar.HTMLCalendar):
        def __init__(self, l_dict, h_dict, yr, mo):
            super().__init__()
            self.l_dict = l_dict
            self.h_dict = h_dict
            self.yr = yr
            self.mo = mo

        def formatday(self, day, weekday):
            if day == 0: return '<td style="background-color:#fafafa; border:1px solid #eee;">&nbsp;</td>'
            date_str = f"{self.yr}-{self.mo:02d}-{day:02d}"
            
            cell_html = f"<strong style='font-size: 14px; color: #333;'>{day}</strong>"
            bg_color = "white"
            
            holiday_name = self.h_dict.get(date_str)
            if holiday_name:
                cell_html += f"<br><span style='font-size:9px; color:#0055cc; font-weight:bold; line-height:1;'>{holiday_name}</span>"
            
            if date_str in self.l_dict:
                bg_color = "#ffcccc"
                names = "<br>".join([f"<span style='font-size:10px; color:white; background:#cc0000; padding:2px 4px; border-radius:3px; display:inline-block; margin-top:2px; white-space:nowrap;'>{n}</span>" for n in self.l_dict[date_str]])
                cell_html += names
                
            return f'<td style="background-color:{bg_color}; border:1px solid #ddd; vertical-align:top; height:70px; width:14%; padding:3px;">{cell_html}</td>'

        def formatmonthname(self, theyear, themonth, withyear=True):
            month_name = calendar.month_name[themonth]
            return f'<tr><th colspan="7" style="background-color:#333; color:white; font-size:16px; padding:8px; border-radius: 5px 5px 0 0;">{month_name} {theyear}</th></tr>'
            
        def formatweekheader(self):
            s = ''.join(f'<th style="background-color:#f0f2f6; color:#333; padding:5px; font-size:12px; border:1px solid #ddd;">{wk}</th>' for wk in ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"])
            return f'<tr>{s}</tr>'

    for row in range(4):
        cols = st.columns(3)
        for col_idx in range(3):
            offset = row * 3 + col_idx
            calc_month = cur_month + offset
            calc_year = cur_year
            while calc_month > 12:
                calc_month -= 12
                calc_year += 1

            cal = LeaveCalendar(leave_dict, us_holidays, calc_year, calc_month)
            cal_html = cal.formatmonth(calc_year, calc_month)
            
            cal_html = cal_html.replace('<table border="0" cellpadding="0" cellspacing="0" class="month">', '<table style="width:100%; border-collapse:collapse; text-align:center; font-family:sans-serif; margin-bottom: 20px; box-shadow: 0px 4px 10px rgba(0,0,0,0.1);">')
            
            with cols[col_idx]:
                st.markdown(cal_html, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Manage Raw Requests")
    if leave_requests:
        df_leave = pd.DataFrame(leave_requests).sort_values(by="start")
        st.dataframe(df_leave, use_container_width=True, hide_index=True)
        
        if st.button("🗑️ Clear All Requests (Admin)"):
            save_json_db(LEAVE_FILE, [])
            st.rerun()
    else: st.info("No time off requested yet.")